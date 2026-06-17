import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    with open('cruds_content.txt', 'w', encoding='utf-8') as f:
        for i, para in enumerate(doc.paragraphs):
            f.write(f"Para {i}: {para.text}\n")
        
        f.write("\n\n--- TABLES ---\n\n")
        for t_idx, table in enumerate(doc.tables):
            f.write(f"Table {t_idx}:\n")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    f.write(f"  Row {r_idx}, Col {c_idx}: {cell.text}\n")

if __name__ == "__main__":
    read_docx("CRUDs.docx")
